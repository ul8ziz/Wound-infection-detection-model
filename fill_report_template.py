"""
Fill the intern report template (отчет-практиканта-шаблон.docx) with wound detection project content.
Output: отчет-практиканта-заполненный.docx
"""
from pathlib import Path

from docx import Document

# Paths
ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "отчет-практиканта-шаблон.docx"
OUTPUT = ROOT / "отчет-практиканта-заполненный.docx"

# Content
THESIS_TITLE = "Детекция и сегментация послеоперационных ран с использованием Mask R-CNN"

WORK_TABLE_ROWS = [
    ("Установочное занятие.", "Инструктаж по безопасности труда и правилам пожарной безопасности при выполнении лабораторных и практических работ. Обсуждение задания на практику. Тема: детекция и сегментация послеоперационных ран с использованием Mask R-CNN.", "07.02.2026"),
    ("Подбор материалов", "Подбор материалов для написания ВКР. Анализ литературы по Mask R-CNN, сегментации медицинских изображений, детекции ран.", "03.09.-30.09.2026"),
    ("Подготовка данных", "Изучение структуры COCO, аннотаций CVAT. Скрипты build_wound_focus_dataset, build_wound_only_dataset. Формирование wound_focus_clean.", "01.10.-23.10.2026"),
    ("Практическая часть", "Разработка пайплайна обучения. Настройка Mask R-CNN ResNet-50-FPN, pipeline_utils, augmentation_strategy.", "24.10.-06.11.2026"),
    ("Численный эксперимент", "Обучение модели, валидация датасета. Метрики COCO: bbox_AP50, segm_AP50, combined_AP50.", "07.11.-27.11.2026"),
    ("Оптимизация и анализ", "Расчёт площади раны в см² (pixels_per_cm). Интерпретация результатов, анализ разрыва валидация–тест.", "28.11.-05.12.2026"),
    ("Написание отчёта", "Описание методов, архитектуры модели, пайплайна данных, результатов обучения. Выводы и ограничения.", "06.12-12.12.2026"),
    ("Подготовка отчёта и дневника", "Собеседование с научным руководителем и руководителем практики по содержанию отчёта и дневника по практике.", "25.12-30.12.2026"),
    ("Сдача отчёта и дневника", "Сдача отчёта и дневника по практике руководителю практики в ТУИС.", "07.03.2026"),
]


def replace_in_paragraph(paragraph: "Paragraph", old: str, new: str) -> bool:
    """Replace text in a paragraph. Returns True if replacement was made."""
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Text may be split across runs
    full = "".join(r.text for r in paragraph.runs)
    if old in full:
        paragraph.clear()
        paragraph.add_run(full.replace(old, new))
        return True
    return False


def replace_in_document(doc: Document, old: str, new: str) -> int:
    """Replace all occurrences of old with new in document. Returns count."""
    count = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, old, new):
                        count += 1
    return count


def fill_work_table(doc: Document) -> None:
    """Fill the work sequence table. Table has header row + data rows."""
    for table in doc.tables:
        # Find table with "Работы и мероприятия" header
        header_text = "".join(c.text for c in table.rows[0].cells).lower()
        if "работы" in header_text and "мероприятия" in header_text:
            # Header: № п/п | Работы и мероприятия | Пояснение | Сроки выполнения
            # Data rows start at index 1
            for i, (work, desc, dates) in enumerate(WORK_TABLE_ROWS):
                row_idx = i + 1
                if row_idx >= len(table.rows):
                    table.add_row()
                    row = table.rows[-1]
                else:
                    row = table.rows[row_idx]
                cells = row.cells
                if len(cells) >= 4:
                    cells[0].text = str(i + 1)
                    cells[1].text = work
                    cells[2].text = desc
                    cells[3].text = dates
            return
    # Fallback: try first table with 4 columns
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.rows[0].cells) >= 4:
            for i, (work, desc, dates) in enumerate(WORK_TABLE_ROWS):
                row_idx = i + 1
                if row_idx >= len(table.rows):
                    table.add_row()
                    row = table.rows[-1]
                else:
                    row = table.rows[row_idx]
                cells = row.cells
                if len(cells) >= 4:
                    cells[0].text = str(i + 1)
                    cells[1].text = work
                    cells[2].text = desc
                    cells[3].text = dates
            return


def add_main_part_content(doc: Document) -> None:
    """Add main part content after 'Основная часть' heading, before 'Заключение'."""
    main_content = [
        "В рамках производственной практики была реализована система детекции и сегментации послеоперационных ран на клинических фотографиях с использованием архитектуры Mask R-CNN.",
        "",
        "Датасет. Использован датасет wound_focus_clean в формате COCO с полигональными аннотациями из CVAT. Разделение: train 257, val 57, test 55 изображений.",
        "",
        "Модель. Mask R-CNN ResNet-50-FPN с num_classes=2 (фон и рана), размер входа 512×512 пикселей. Обучение с аугментациями Albumentations, ранняя остановка по метрике combined_AP50.",
        "",
        "Результаты. Лучшие метрики на валидации: combined_AP50 0,43, bbox_AP50 0,52, segm_AP50 0,35. На тесте: combined_AP50 0,33. Реализован расчёт площади раны в см² с использованием параметра pixels_per_cm при отсутствии маркера.",
        "",
        "Ограничения. Проект носит исследовательский характер и не предназначен для клинического применения без дополнительной валидации. Качество аннотаций ограничивает возможность надёжной сегментации подклассов (фибрин, грануляции, некроз и др.).",
    ]
    for p in doc.paragraphs:
        if p.text.strip() == "Заключение":
            ref = p._element
            for line in reversed(main_content):
                new_p = doc.add_paragraph(line)
                new_p._element.getparent().remove(new_p._element)
                ref.addprevious(new_p._element)
                ref = new_p._element
            return


def adapt_conclusion(doc: Document) -> None:
    """Replace generic conclusion text with wound-detection-specific content."""
    old_iot = "Выполненный во время проведения производственной практики обзор публикаций научных изданий как по теме Интернет вещей, так и по теме анализа характеристик интерференции, возникающей при прямом взаимодействии устройств, позволит мне обосновать актуальность выбранной темы, а также более полно раскрыть"
    new_iot = "Выполненный во время проведения производственной практики обзор публикаций по Mask R-CNN, сегментации медицинских изображений и детекции ран, а также реализованный пайплайн обучения и оценки, позволят обосновать актуальность выбранной темы ВКР и более полно раскрыть практическую часть выпускной работы."
    replace_in_document(doc, old_iot, new_iot)


def add_abbreviations(doc: Document) -> None:
    """Add abbreviations to Список сокращений."""
    abbrevs_ru = [
        ("ВКР", "выпускная квалификационная работа"),
        ("РУДН", "Российский университет дружбы народов"),
        ("ТУИС", "телематическая учебно-информационная система"),
    ]
    abbrevs_en = [
        ("COCO", "Common Objects in Context"),
        ("CVAT", "Computer Vision Annotation Tool"),
        ("FPN", "Feature Pyramid Network"),
        ("AP50", "Average Precision at IoU 0.5"),
        ("Mask R-CNN", "Mask Region-based Convolutional Neural Network"),
        ("ResNet", "Residual Network"),
    ]
    # Find "Англоязычные сокращения" and add entries after
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Англоязычные сокращения" in cell.text:
                    # Add abbreviations as new paragraphs in this cell or next
                    pass
    # Use replace to add content in Список сокращений section
    # Simpler: replace "Англоязычные сокращения" with full list
    old_en = "Англоязычные сокращения"
    new_en = "Англоязычные сокращения: COCO — Common Objects in Context; CVAT — Computer Vision Annotation Tool; FPN — Feature Pyramid Network; AP50 — Average Precision at IoU 0.5; Mask R-CNN — Mask Region-based Convolutional Neural Network; ResNet — Residual Network."
    replace_in_document(doc, old_en, new_en)


def add_sources(doc: Document) -> None:
    """Add list of sources after 'Список источников' section heading."""
    sources = [
        "1. He K., Gkioxari G., Dollár P., Girshick R. Mask R-CNN // IEEE International Conference on Computer Vision (ICCV). 2017. P. 2961–2969.",
        "2. Lin T.-Y. et al. Microsoft COCO: Common Objects in Context // ECCV. 2014. P. 740–755.",
        "3. PyTorch: An open source machine learning framework. https://pytorch.org/",
        "4. Albumentations: Fast image augmentation library. https://albumentations.ai/",
        "5. CVAT: Computer Vision Annotation Tool. https://www.cvat.ai/",
    ]
    # Find last "Список источников" (section heading) and insert sources after it
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]
        if "Список источников" in p.text and "Оглавление" not in p.text:
            ref = p._element
            for s in sources:
                new_p = doc.add_paragraph(s)
                elem = new_p._element
                elem.getparent().remove(elem)
                ref.addnext(elem)
                ref = elem
            return


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")

    doc = Document(str(TEMPLATE))

    # 1. Replace thesis title placeholder
    replace_in_document(doc, "НАЗВАНИЕ ВЫПУСКНОЙ РАБОТЫ", THESIS_TITLE)

    # 2. Fill work sequence table
    fill_work_table(doc)

    # 3. Adapt conclusion (replace IoT/interference text)
    adapt_conclusion(doc)

    # 4. Add abbreviations
    add_abbreviations(doc)

    # 5. Add main part content (insert before Заключение)
    add_main_part_content(doc)

    # 6. Add sources
    add_sources(doc)

    doc.save(str(OUTPUT))
    # Avoid printing Cyrillic to console on Windows
    with open(ROOT / "fill_report_log.txt", "w", encoding="utf-8") as f:
        f.write(f"Saved: {OUTPUT}\n")


if __name__ == "__main__":
    main()
